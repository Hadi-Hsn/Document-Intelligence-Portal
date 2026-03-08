"""
Text Extraction Module
Handles extraction of text from PDF, DOCX, and image files.
"""

import io
import os
import tempfile
from pathlib import Path
from typing import Optional

from PIL import Image
from docx import Document
import pdfplumber
import pytesseract


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Route file to the appropriate extractor based on extension.
    Returns extracted text as a string.
    """
    ext = Path(filename).suffix.lower()

    if ext in (".pdf",):
        return extract_from_pdf(file_bytes)
    elif ext in (".doc", ".docx"):
        return extract_from_docx(file_bytes)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"):
        return extract_from_image(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF files.
    - Uses pdfplumber for embedded text.
    - Falls back to OCR for scanned/image-based pages.
    """
    all_text = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            # Try extracting embedded text first
            text = page.extract_text() or ""

            # If page has very little text, it might be a scanned image — try OCR
            if len(text.strip()) < 50:
                ocr_text = _ocr_pdf_page(page)
                if ocr_text and len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text

            if text.strip():
                all_text.append(f"--- Page {i + 1} ---\n{text.strip()}")

    return "\n\n".join(all_text)


def _ocr_pdf_page(page) -> str:
    """OCR a single pdfplumber page by converting it to an image."""
    try:
        # Convert page to image
        img = page.to_image(resolution=300)
        pil_image = img.original

        # Run OCR
        text = pytesseract.image_to_string(pil_image)
        return text
    except Exception:
        return ""


def extract_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from Word (.docx) files.
    Extracts paragraphs, tables, and core metadata.
    """
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Extract metadata
    props = doc.core_properties
    meta_lines = []
    if props.title:
        meta_lines.append(f"Title: {props.title}")
    if props.author:
        meta_lines.append(f"Author: {props.author}")
    if props.created:
        meta_lines.append(f"Created: {props.created}")
    if props.modified:
        meta_lines.append(f"Modified: {props.modified}")
    if props.subject:
        meta_lines.append(f"Subject: {props.subject}")

    if meta_lines:
        parts.append("--- Document Metadata ---\n" + "\n".join(meta_lines))

    # Extract paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    if paragraphs:
        parts.append("--- Document Text ---\n" + "\n".join(paragraphs))

    # Extract tables
    table_texts = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            table_texts.append(f"Table {t_idx + 1}:\n" + "\n".join(rows))

    if table_texts:
        parts.append("--- Tables ---\n" + "\n\n".join(table_texts))

    return "\n\n".join(parts)


def extract_from_image(file_bytes: bytes) -> str:
    """Extract text from image files using OCR."""
    image = Image.open(io.BytesIO(file_bytes))

    # Convert to RGB if necessary (e.g., RGBA or palette images)
    if image.mode not in ("L", "RGB"):
        image = image.convert("RGB")

    text = pytesseract.image_to_string(image)
    return text.strip()
